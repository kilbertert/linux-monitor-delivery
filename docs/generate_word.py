#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
from docx import Document
from docx.shared import Pt, Mm, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 页边距设置 (单位: mm)
TOP_MARGIN = 30
BOTTOM_MARGIN = 25
LEFT_MARGIN = 30
RIGHT_MARGIN = 20

def set_page_margins(section):
    """设置页边距"""
    section.top_margin = Mm(TOP_MARGIN)
    section.bottom_margin = Mm(BOTTOM_MARGIN)
    section.left_margin = Mm(LEFT_MARGIN)
    section.right_margin = Mm(RIGHT_MARGIN)

def read_markdown_content(filepath):
    """读取Markdown文件内容"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def parse_markdown_to_paragraphs(content):
    """解析Markdown内容为段落列表"""
    lines = content.split('\n')
    paragraphs = []
    current_paragraph = []
    in_code_block = False
    table_lines = []
    in_table = False
    
    for line in lines:
        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(current_paragraph)
                if code_text:
                    paragraphs.append(('code', code_text))
                current_paragraph = []
                in_code_block = False
            else:
                if current_paragraph:
                    paragraphs.append(('text', '\n'.join(current_paragraph)))
                    current_paragraph = []
                in_code_block = True
            continue
        
        if in_code_block:
            current_paragraph.append(line)
            continue
        
        # 处理表格
        if line.strip().startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue
        elif in_table:
            paragraphs.append(('table', table_lines[:]))
            table_lines = []
            in_table = False
        
        # 处理标题和正文
        if line.startswith('# '):
            if current_paragraph:
                paragraphs.append(('text', '\n'.join(current_paragraph)))
                current_paragraph = []
            paragraphs.append(('title', line[2:].strip()))
        elif line.startswith('## '):
            if current_paragraph:
                paragraphs.append(('text', '\n'.join(current_paragraph)))
                current_paragraph = []
            paragraphs.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            if current_paragraph:
                paragraphs.append(('text', '\n'.join(current_paragraph)))
                current_paragraph = []
            paragraphs.append(('h3', line[4:].strip()))
        elif line.strip():
            current_paragraph.append(line)
        elif current_paragraph:
            paragraphs.append(('text', '\n'.join(current_paragraph)))
            current_paragraph = []
    
    # 处理最后的段落
    if current_paragraph:
        if in_code_block:
            paragraphs.append(('code', '\n'.join(current_paragraph)))
        else:
            paragraphs.append(('text', '\n'.join(current_paragraph)))
    
    if in_table and table_lines:
        paragraphs.append(('table', table_lines))
    
    return paragraphs

def parse_table_markdown(table_lines):
    """解析表格Markdown"""
    if len(table_lines) < 2:
        return [], []
    
    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
    
    rows = []
    for line in table_lines[2:]:
        if line.strip().startswith('|') and '|' in line:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)
    
    return headers, rows

def format_document(filepath, output_path):
    """格式化文档"""
    doc = Document()
    section = doc.sections[0]
    set_page_margins(section)
    
    content = read_markdown_content(filepath)
    paragraphs = parse_markdown_to_paragraphs(content)
    
    for para in paragraphs:
        para_type, para_content = para
        
        if para_type == 'title':
            p = doc.add_paragraph()
            run = p.add_run(para_content)
            run.font.size = Pt(22)
            run.font.name = '黑体'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            
        elif para_type == 'h2':
            p = doc.add_paragraph()
            run = p.add_run(para_content)
            run.font.size = Pt(16)
            run.font.name = '黑体'
            run.font.bold = True
            p.paragraph_format.line_spacing = 1.5
            
        elif para_type == 'h3':
            p = doc.add_paragraph()
            run = p.add_run(para_content)
            run.font.size = Pt(14)
            run.font.name = '黑体'
            run.font.bold = True
            p.paragraph_format.line_spacing = 1.5
            
        elif para_type == 'text':
            p = doc.add_paragraph()
            run = p.add_run(para_content)
            run.font.size = Pt(12)
            run.font.name = '宋体'
            p.paragraph_format.line_spacing = 1.5
            
        elif para_type == 'code':
            p = doc.add_paragraph()
            run = p.add_run(para_content)
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.line_spacing = 1.0
            
        elif para_type == 'table':
            headers, rows = parse_table_markdown(para_content)
            if headers and rows:
                table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                for i, header in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = header
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                    cell.paragraphs[0].runs[0].font.name = '宋体'
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(headers):
                            cell = table.rows[row_idx+1].cells[col_idx]
                            cell.text = cell_data
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
    
    doc.save(output_path)
    print(f"Word文档已生成: {output_path}")

if __name__ == '__main__':
    input_file = '/home/rl/.openclaw/workspace/delivery/linux-monitor/docs/毕业设计论文.md'
    output_file = '/home/rl/.openclaw/workspace/delivery/linux-monitor/docs/毕业设计论文.docx'
    format_document(input_file, output_file)
