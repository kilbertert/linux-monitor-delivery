#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

# 页边距设置 (单位: mm)
TOP_MARGIN = 30
BOTTOM_MARGIN = 25
LEFT_MARGIN = 30
RIGHT_MARGIN = 20

def set_page_margins(section):
    section.top_margin = Mm(TOP_MARGIN)
    section.bottom_margin = Mm(BOTTOM_MARGIN)
    section.left_margin = Mm(LEFT_MARGIN)
    section.right_margin = Mm(RIGHT_MARGIN)

def read_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def parse_content(content):
    lines = content.split('\n')
    blocks = []
    current_text = []
    in_code = False
    in_table = False
    table_data = []
    
    for line in lines:
        if line.startswith('```'):
            if in_code:
                blocks.append(('code', '\n'.join(current_text)))
                current_text = []
                in_code = False
            else:
                if current_text:
                    blocks.append(('text', '\n'.join(current_text)))
                    current_text = []
                in_code = True
            continue
        
        if in_code:
            current_text.append(line)
            continue
        
        if line.strip().startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_data = []
            table_data.append(line)
            continue
        elif in_table:
            blocks.append(('table', table_data[:]))
            table_data = []
            in_table = False
        
        if line.startswith('# '):
            if current_text:
                blocks.append(('text', '\n'.join(current_text)))
                current_text = []
            blocks.append(('title', line[2:].strip()))
        elif line.startswith('## '):
            if current_text:
                blocks.append(('text', '\n'.join(current_text)))
                current_text = []
            blocks.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            if current_text:
                blocks.append(('text', '\n'.join(current_text)))
                current_text = []
            blocks.append(('h3', line[4:].strip()))
        elif line.strip():
            current_text.append(line)
        elif current_text:
            blocks.append(('text', '\n'.join(current_text)))
            current_text = []
    
    if current_text:
        if in_code:
            blocks.append(('code', '\n'.join(current_text)))
        else:
            blocks.append(('text', '\n'.join(current_text)))
    
    if in_table and table_data:
        blocks.append(('table', table_data))
    
    return blocks

def parse_table(table_lines):
    if len(table_lines) < 2:
        return [], []
    headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
    rows = []
    for line in table_lines[2:]:
        if line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                rows.append(cells)
    return headers, rows

def create_doc(input_path, output_path):
    doc = Document()
    section = doc.sections[0]
    set_page_margins(section)
    
    content = read_file(input_path)
    blocks = parse_content(content)
    
    # 定义样式
    for block in blocks:
        btype, bcontent = block
        
        if btype == 'title':
            p = doc.add_paragraph()
            run = p.add_run(bcontent)
            run.font.size = Pt(22)
            run.font.name = '黑体'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            
        elif btype == 'h2':
            p = doc.add_paragraph()
            run = p.add_run(bcontent)
            run.font.size = Pt(16)
            run.font.name = '黑体'
            run.font.bold = True
            p.paragraph_format.line_spacing = 1.5
            
        elif btype == 'h3':
            p = doc.add_paragraph()
            run = p.add_run(bcontent)
            run.font.size = Pt(14)
            run.font.name = '黑体'
            run.font.bold = True
            p.paragraph_format.line_spacing = 1.5
            
        elif btype == 'text':
            # 处理粗体文本
            segments = re.split(r'(\*\*[^*]+\*\*)', bcontent)
            if len(segments) == 1:
                p = doc.add_paragraph()
                run = p.add_run(bcontent)
                run.font.size = Pt(12)
                run.font.name = '宋体'
            else:
                p = doc.add_paragraph()
                for seg in segments:
                    if seg.startswith('**') and seg.endswith('**'):
                        run = p.add_run(seg.replace('**', ''))
                        run.font.size = Pt(12)
                        run.font.name = '宋体'
                        run.font.bold = True
                    elif seg:
                        run = p.add_run(seg)
                        run.font.size = Pt(12)
                        run.font.name = '宋体'
            p.paragraph_format.line_spacing = 1.5
            
        elif btype == 'code':
            p = doc.add_paragraph()
            run = p.add_run(bcontent)
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.line_spacing = 1.0
            
        elif btype == 'table':
            headers, rows = parse_table(bcontent)
            if headers and rows:
                table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                    cell.paragraphs[0].runs[0].font.name = '宋体'
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(headers):
                            table.rows[ri+1].cells[ci].text = cell
                            table.rows[ri+1].cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
                            table.rows[ri+1].cells[ci].paragraphs[0].runs[0].font.name = 'Times New Roman'
    
    doc.save(output_path)
    print(f"Word文档已生成: {output_path}")

if __name__ == '__main__':
    create_doc(
        '/home/rl/.openclaw/workspace/delivery/linux-monitor/docs/毕业设计论文.md',
        '/home/rl/.openclaw/workspace/delivery/linux-monitor/docs/毕业设计论文.docx'
    )
