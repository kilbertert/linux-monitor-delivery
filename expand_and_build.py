import os
import re
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, ascii_name='Times New Roman', eastasia_name='宋体', size=12, bold=False):
    run.font.name = ascii_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), eastasia_name)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "1"
    fldChar2.append(t)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def expand_markdown(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Expand "关键技术实现"
    if "### 5.1.1 CPU 采集实现" in content and "**核心代码如下**" not in content:
        try:
            with open('backend/monitor/collector.py', 'r', encoding='utf-8') as f_col:
                col_code = f_col.read()
                match = re.search(r'def get_cpu.*?return\s+\{.*?\}', col_code, re.DOTALL)
                if match:
                    cpu_code = match.group(0)
                    expansion = f"\n\n**核心代码如下**：\n```python\n{cpu_code}\n```\n代码解析：首先通过 `_warmup` 预热，然后使用 `psutil.cpu_percent(percpu=True)` 获取各核心的占用率，这对于多核服务器性能分析至关重要。\n"
                    content = content.replace("具体代码实现如下：首先使用 psutil.cpu_percent(interval=None) 进行初始化预热，然后等待 0.1 秒后再次调用获取各核心使用率，最后计算总体使用率。返回数据包含 percent（总体使用率）、per_core（各核心使用率列表）和 count（CPU 核心数）。", 
                                              "具体代码实现如下：首先使用 psutil.cpu_percent(interval=None) 进行初始化预热，然后等待 0.1 秒后再次调用获取各核心使用率，最后计算总体使用率。返回数据包含 percent（总体使用率）、per_core（各核心使用率列表）和 count（CPU 核心数）。" + expansion)
        except Exception:
            pass

    # Expand WebSocket
    if "### 5.3 WebSocket 实时推送实现" in content and "**核心代码如下**" not in content:
        try:
            with open('backend/api/websocket.py', 'r', encoding='utf-8') as f_ws:
                ws_code = f_ws.read()
                match = re.search(r'class ConnectionManager:.*?async def broadcast.*?:\n.*?pass', ws_code, re.DOTALL)
                if not match:
                    match = re.search(r'class ConnectionManager:.*?(?=class|\Z)', ws_code, re.DOTALL)
                if match:
                    code_str = match.group(0)[:800] # limit length
                    expansion = f"\n\n**核心代码如下**：\n```python\n{code_str}\n```\n代码解析：ConnectionManager 类负责管理所有活跃的 WebSocket 连接。当客户端建立连接时，会将其加入活跃列表；当客户端断开时，会从列表中移除。这种集中管理方式使得后端可以方便地向所有在线客户端广播最新的系统指标数据。\n"
                    content = content.replace("客户端可以通过发送 {\"action\": \"subscribe\", \"metrics\": [\"cpu\", \"memory\"], \"interval\": 1} 消息订阅指标，通过发送 {\"action\": \"unsubscribe\"} 取消订阅。后端定时采集数据并推送到所有订阅的客户端。", 
                                              "客户端可以通过发送 {\"action\": \"subscribe\", \"metrics\": [\"cpu\", \"memory\"], \"interval\": 1} 消息订阅指标，通过发送 {\"action\": \"unsubscribe\"} 取消订阅。后端定时采集数据并推送到所有订阅的客户端。" + expansion)
        except Exception:
            pass

    # Expand Deployment
    if "### 7.4 Nginx 反向代理配置" in content and "server {" not in content:
        nginx_conf = """
```nginx
server {
    listen 80;
    server_name monitor.example.com;

    # 前端静态资源
    location / {
        root /var/www/monitor/dist;
        index index.html;
        try_files $uri $uri/ /index.html;
    }

    # REST API 代理
    location /api/ {
        proxy_pass http://127.0.0.1:8000/;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
    }

    # WebSocket 代理
    location /ws/ {
        proxy_pass http://127.0.0.1:8000/ws/;
        proxy_http_version 1.1;
        proxy_set_header Upgrade $http_upgrade;
        proxy_set_header Connection "Upgrade";
        proxy_set_header Host $host;
        proxy_read_timeout 3600s;
        proxy_send_timeout 3600s;
    }
}
```
配置详解：在上述 Nginx 配置中，前端静态资源通过 `root` 指令指定路径，并使用 `try_files` 支持 Vue Router 的 History 模式。API 请求代理至后端的 8000 端口，并传递客户端的真实 IP。WebSocket 代理则特别配置了 `Upgrade` 和 `Connection` 头部，以维持长连接的稳定性，同时设置了较长的超时时间防止连接意外断开。
"""
        content = content.replace("配置文件示例：\n\n前端静态文件部署在 /var/www/monitor/dist，API 代理到 http://localhost:8000，WebSocket 代理需要设置 Upgrade 和 Connection 头。", 
                                  "配置文件示例：\n" + nginx_conf)

    # Expand Database
    if "### 4.3 数据库设计" in content and "**核心建表语句**" not in content:
        try:
            with open('backend/db/database.py', 'r', encoding='utf-8') as f_db:
                db_code = f_db.read()
                match = re.search(r'def init_db.*?cursor\.executescript.*?(\"\"\".*?\"\"\")', db_code, re.DOTALL)
                if match:
                    schema = match.group(1)
                    expansion = f"\n\n**核心建表语句**：\n```sql\n{schema}\n```\n通过合理设置主键和索引，保证了时间序列数据的快速查询。其中 timestamp 字段采用整数类型存储 Unix 时间戳，极大提高了范围查询的效率。\n"
                    content = content.replace("各表在 timestamp 字段上建立索引，以加快历史数据查询速度。", 
                                              "各表在 timestamp 字段上建立索引，以加快历史数据查询速度。" + expansion)
        except Exception:
            pass

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(content)
        
    return content

def create_word_document(md_content, out_path):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Mm(30)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(30)
        section.right_margin = Mm(20)

    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style_normal.paragraph_format.line_spacing = 1.5

    lines = md_content.split('\n')
    
    # First, collect headers for TOC
    headers = []
    for line in lines:
        line = line.strip()
        if line.startswith('## ') and '摘要' not in line and 'ABSTRACT' not in line and '毕业设计' not in line:
            headers.append((1, line[3:]))
        elif line.startswith('### '):
            headers.append((2, line[4:]))
        elif line.startswith('#### '):
            headers.append((3, line[5:]))
            
    # Now build the document
    in_code_block = False
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(line)
            set_font(run, ascii_name='Consolas', eastasia_name='宋体', size=10)
            continue
            
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(line[2:])
            set_font(run, ascii_name='Times New Roman', eastasia_name='黑体', size=16, bold=True)
            
        elif line.startswith('## 摘要') or line.startswith('## ABSTRACT'):
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(line[3:].strip())
            set_font(run, ascii_name='Times New Roman', eastasia_name='黑体', size=16, bold=True)
            
            # Insert TOC after ABSTRACT
            if 'ABSTRACT' in line:
                # Add page break
                run.add_break(WD_BREAK.PAGE)
                
                # Add TOC Title
                p_toc = doc.add_paragraph()
                p_toc.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p_toc.paragraph_format.line_spacing = 1.5
                run_toc = p_toc.add_run("目录")
                set_font(run_toc, ascii_name='Times New Roman', eastasia_name='黑体', size=16, bold=True)
                
                # Add TOC Items
                for level, text in headers:
                    p_item = doc.add_paragraph()
                    p_item.paragraph_format.line_spacing = 1.5
                    p_item.paragraph_format.left_indent = Mm((level-1) * 5)
                    run_item = p_item.add_run(text)
                    if level == 1:
                        set_font(run_item, ascii_name='Times New Roman', eastasia_name='黑体', size=12, bold=False) # 第一级标题小四号黑体
                    else:
                        set_font(run_item, ascii_name='Times New Roman', eastasia_name='宋体', size=12, bold=False) # 其余小四号宋体
                
                # Add page break after TOC
                doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
                
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.style = doc.styles['Heading 1']
            p.text = ''
            run = p.add_run(line[3:])
            set_font(run, ascii_name='Times New Roman', eastasia_name='黑体', size=18, bold=True)
            
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.style = doc.styles['Heading 2']
            p.text = ''
            run = p.add_run(line[4:])
            set_font(run, ascii_name='Times New Roman', eastasia_name='黑体', size=16, bold=True)
            
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.style = doc.styles['Heading 3']
            p.text = ''
            run = p.add_run(line[5:])
            set_font(run, ascii_name='Times New Roman', eastasia_name='黑体', size=14, bold=True)
            
        elif line.startswith('**关键词**') or line.startswith('**Keywords**'):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            
            # 关键词：四号黑体加粗
            run_title = p.add_run(line.split('**')[1] + line.split('**')[2][0])
            set_font(run_title, ascii_name='Times New Roman', eastasia_name='黑体', size=14, bold=True)
            
            # The rest of the keywords
            rest = line.split('**')[2][1:] if len(line.split('**')) > 2 else ''
            if rest:
                run_rest = p.add_run(rest)
                set_font(run_rest, ascii_name='Times New Roman', eastasia_name='宋体', size=12)
                
        else:
            clean_line = line.replace('**', '')
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(clean_line)
            set_font(run, ascii_name='Times New Roman', eastasia_name='宋体', size=12)
            
    # Add page numbers
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = p.add_run()
        set_font(run, ascii_name='Times New Roman', eastasia_name='Times New Roman', size=12)
        add_page_number(run)

    doc.save(out_path)
    print("Document saved.")

if __name__ == "__main__":
    md_file = 'docs/毕业设计论文.md'
    out_file = 'docs/毕业设计论文.docx'
    content = expand_markdown(md_file)
    create_word_document(content, out_file)
